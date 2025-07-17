from pathlib import Path
from llama_parse import LlamaParse
import docx
import pypandoc
import shutil
import os
import logging
from langchain_core.documents import Document
import config
from db.weaviateDB import connect_to_weaviate
import utils.utils as utils
logger = logging.getLogger(__name__)

from rag_components import create_weaviate_schema_if_not_exists, ingest_chunks_with_native_batching
from utils.process_data import hierarchical_split_law_document,extract_document_metadata,clean_document_text,infer_field, infer_entity_type, filter_and_serialize_complex_metadata

def convert_to_text_content(source_path: str) -> str:
    source_file = Path(source_path)
    file_extension = source_file.suffix.lower()
    logger.info(f"Extracting content from: {source_file.name}")
    content = ""
    if file_extension == ".pdf":
        parser = LlamaParse( api_key=config.LLAMA_CLOUD_API_KEY,
                    result_type="text",
                    verbose=True, # Giữ verbose để theo dõi
                    language="vi")
        documents = parser.load_data([str(source_file)])
        if documents: content = documents[0].text
    elif file_extension == ".docx":
        doc = docx.Document(source_path)
        content = '\n'.join([para.text for para in doc.paragraphs])
    elif file_extension == ".doc":
        content = pypandoc.convert_file(source_path, 'plain', format='doc')
    else:
        raise ValueError(f"Unsupported file format: {file_extension}")
    if not content.strip():
        raise ValueError("Extracted content is empty.")
    logger.info(f"✅ Successfully extracted content from {source_file.name}.")
    return content

def full_process_and_ingest_pipeline(filepath: str, file_hash: str, embedding_model):
    filename = os.path.basename(filepath)
    logger.info(f"BACKGROUND TASK: Starting full pipeline for: {filename} (Hash: {file_hash[:10]}...)")
    weaviate_client = None
    try:
        raw_content = convert_to_text_content(filepath)

        doc_metadata = extract_document_metadata(raw_content, filename)
        doc_metadata["source"] = filename
        cleaned_content = clean_document_text(raw_content)
        doc_metadata["field"] = infer_field(cleaned_content, doc_metadata.get("ten_van_ban"))
        doc_metadata["entity_type"] = infer_entity_type(cleaned_content, doc_metadata.get("field", ""))

        doc_to_split = Document(page_content=cleaned_content, metadata=doc_metadata)
        chunks_from_file = hierarchical_split_law_document(doc_to_split)

        if not chunks_from_file:
            raise ValueError("File did not yield any chunks after processing.")

        processed_chunks = filter_and_serialize_complex_metadata(chunks_from_file)

        weaviate_client = connect_to_weaviate()
        embeddings_model = embedding_model
        collection_name = config.WEAVIATE_COLLECTION_NAME
        create_weaviate_schema_if_not_exists(weaviate_client, collection_name)

        ingest_chunks_with_native_batching(weaviate_client, collection_name, processed_chunks, embeddings_model)

        utils.log_processed_hash(file_hash)
        logger.info(f"✅ Successfully ingested '{filename}'.")
        shutil.move(filepath, os.path.join(config.PROCESSED_FILES_FOLDER, filename))
        logger.info(f"Moved '{filename}' to processed folder.")
    except Exception as e:
        logger.error(f"❌ FAILED pipeline for '{filename}': {e}", exc_info=True)
        shutil.move(filepath, os.path.join(config.FAILED_FILES_FOLDER, filename))
        logger.info(f"Moved '{filename}' to failed folder.")
    finally:
        if weaviate_client and weaviate_client.is_connected():
            weaviate_client.close()