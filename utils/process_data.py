import re
import os
from typing import List, Dict, Optional, Tuple, Union,Any
import logging
from tqdm import tqdm
import uuid
import json
import time
from langchain_core.documents import Document
from config import LEGAL_DOC_TYPES, MAX_CHUNK_SIZE, CHUNK_OVERLAP, model_process
from docx import Document as DocxDocument
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import JsonOutputParser

import prompt_templete
GOOGLE_API_KEY = os.environ.get("GOOGLE_API_KEY")

# Thiết lập logging
# logging.basicConfig(level=logging.DEBUG, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

#Hàm cuối cùng để chuẩn hóa metadata ngay trước khi ingest vào vector store.
def filter_and_serialize_complex_metadata(documents: List[Document]) -> List[Document]:
    """Hàm cuối cùng để chuẩn hóa metadata ngay trước khi ingest vào vector store."""
    updated_documents = []
    allowed_types = (str, bool, int, float, type(None))
    serialize_keys = ["penalties", "cross_references"]

    for doc in documents:
        filtered_metadata = {}
        for key, value in doc.metadata.items():
            if key in serialize_keys and value: # Chỉ serialize nếu value không rỗng
                try:
                    filtered_metadata[key] = json.dumps(value, ensure_ascii=False, default=str)
                except TypeError:
                    logger.warning(f"Không thể serialize key '{key}' cho doc ID {doc.id}. Chuyển thành string.")
                    filtered_metadata[key] = str(value)
            elif isinstance(value, allowed_types):
                filtered_metadata[key] = value
            elif isinstance(value, list):
                filtered_metadata[key] = json.dumps(value, ensure_ascii=False)
            else:
                filtered_metadata[key] = str(value)
        doc.metadata = filtered_metadata
        updated_documents.append(doc)
    return updated_documents

#Một text splitter đơn giản để chia nhỏ các chunk quá lớn.
class SimpleTextSplitter:
    """Một text splitter đơn giản để chia nhỏ các chunk quá lớn."""
    def __init__(self, chunk_size: int, chunk_overlap: int):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def split_text(self, text: str) -> List[str]:
        if not text: return []
        chunks = []
        start = 0
        while start < len(text):
            end = start + self.chunk_size
            chunks.append(text[start:end])
            start += self.chunk_size - self.chunk_overlap
        return chunks

base_text_splitter = SimpleTextSplitter(chunk_size=MAX_CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP)

def generate_structured_id(doc_so_hieu: Optional[str], structure_path: List[str], filename: str) -> str:
    """
    CẢI TIẾN: Tạo ra một chuỗi UUID v5 nhất quán từ ID có cấu trúc.
    Thêm filename để đảm bảo tính duy nhất.
    """
    # Ưu tiên so_hieu, nhưng fallback về filename để tránh "unknown-document"
    base_id = doc_so_hieu if doc_so_hieu else filename
    safe_base_id = re.sub(r'[/\s\.]', '-', base_id) # Thay thế các ký tự không an toàn
    path_str = '_'.join(filter(None, structure_path))

    # Đảm bảo unique_string_id khác nhau cho mỗi file
    unique_string_id = f"{safe_base_id}_{path_str}"

    generated_uuid = uuid.uuid5(uuid.NAMESPACE_DNS, unique_string_id)
    return str(generated_uuid)

def general_ocr_corrections(text: str) -> str:
    """Sửa các lỗi OCR phổ biến."""
    corrections = {
        "LuatWielnam": "LuatVietnam", "LualVietnam": "LuatVietnam",
        "aflu atvistnarn.vni": "@luatvietnam.vn", "Tien ch van bán luat": "Tiện ích văn bản luật",
        "teeeeokanlbaglueloen": "",
        "Tee===": "", "Tc=e===": "", "nem": "",
        "SN Hntlin sa:": "", "HT:": "", "Hntlin sa:": "",
        r"([a-z])([A-Z])": r"\1 \2",
        "Nghịđịnh": "Nghị định", " điểu ": " điều ", "Chưong": "Chương",
        " điềm ": " điểm ", "khoån": "khoản", "Chínhphủ": "Chính phủ",
        " điềukhỏan": " điều khoản", "LuậtTổ": "Luật Tổ", "LuậtXử": "Luật Xử",
        "LuậtTrật": "Luật Trật", " điềuchỉnh": " điều chỉnh", " cá_nhân": " cá nhân",
        " tổ_chức": " tổ chức", " hành_chính": " hành chính", " giấy_phép": " giấy phép",
        " lái_xe": " lái xe", " Giao_thông": " Giao thông",
        # 'ó': '6', 'Ò': '0', 'ọ': '0', 'l': '1', 'I': '1', 'i': '1',
        # 'Z': '2', 'z': '2', 'B': '8',
    }
    for wrong, right in corrections.items():
        if wrong.startswith(r"([a-z])([A-Z])"):
            text = re.sub(wrong, right, text)
        else:
            text = text.replace(wrong, right)
    text = re.sub(r"\s+", " ", text)
    return text.strip()

### HÀM 1: CLEAN_DOCUMENT_TEXT (Cải tiến) ###
def clean_document_text(raw_text: str) -> str:
    """
    Làm sạch văn bản luật một cách an toàn, giữ lại cấu trúc cột ở phần đầu
    và loại bỏ nhiễu một cách có mục tiêu.
    """
    if not raw_text: return ""
    lines = raw_text.splitlines()

    noise_patterns_to_remove = re.compile(
        r"|".join([
            r"LuatVietnam(?:\.vn)?", r"Tiện ích văn bản luật", r"www\.vanbanluat\.vn",
            r"Hotline:", r"Email:", r"Cơ sở dữ liệu văn bản pháp luật",
            r"Trang \d+\s*/\s*\d+", r"^\s*[=\-_*#]+\s*$", r"^\s*\[\s*Hình\s*ảnh\s*]\s*$",
        ]), re.IGNORECASE
    )
    footer_keywords = ["Nơi nhận:", "TM. CHÍNH PHỦ", "TM. BAN BÍ THƯ", "KT. BỘ TRƯỞNG", "TL. BỘ TRƯỞNG", "CHỦ TỊCH QUỐC HỘI"]

    footer_start_index = len(lines)
    for i, line in enumerate(lines):
        line_upper = line.strip().upper()
        if any(keyword.upper() in line_upper for keyword in footer_keywords):
            if "CHỦ TỊCH" in line_upper and i < len(lines) / 2 and "NƯỚC" in line_upper: continue
            footer_start_index = i
            break

    lines_before_footer = lines[:footer_start_index]

    cleaned_lines = []
    for line in lines_before_footer:
        if noise_patterns_to_remove.search(line): continue
        stripped_line = line.strip()
        if not stripped_line: continue
        cleaned_lines.append(stripped_line)

    text = "\n".join(cleaned_lines)
    text = general_ocr_corrections(text)
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    return text

# Hàm tổng hợp phân tích pháp lý toàn diện
def comprehensive_legal_analysis(raw_text: str, filename: str) -> dict:
    """
    Hàm tổng hợp thay thế 3 hàm riêng lẻ:
    1. extract_document_metadata()
    2. hierarchical_split_law_document() 
    3. extract_legal_structures_and_entities()

    Returns:
        dict: {
            "metadata": {...},
            "hierarchical_structure": [...],
            "success": bool,
            "error": str (nếu có lỗi)
        }
    """
    if not GOOGLE_API_KEY:
        raise Exception(
            "GOOGLE_API_KEY is not set. Please provide your Google API key for comprehensive legal analysis.")

    try:

        llm = ChatGoogleGenerativeAI(
            model=model_process,
            google_api_key=GOOGLE_API_KEY,
            temperature=0.0,
        )

        # Sử dụng prompt tổng hợp
        comprehensive_prompt = ChatPromptTemplate.from_template(
            prompt_templete.COMPREHENSIVE_LEGAL_ANALYSIS_PROMPT
        )

        chain = comprehensive_prompt | llm | JsonOutputParser()
        result = chain.invoke({"raw_text": raw_text})
        # for attempt in range(3):
        #     try:
        #         result = chain.invoke({"raw_text": raw_text})
        #         break
        #     except Exception as e:
        #         if "429" in str(e) or "ResourceExhausted" in str(e):
        #             wait_time = (2 ** attempt) * 10
        #             logger.warning(f"⏳ Quá giới hạn API, đợi {wait_time}s rồi thử lại...")
        #             time.sleep(wait_time)
        #         else:
        #             raise e

        # Chuyển đổi key tiếng Việt sang tiếng Anh cho metadata
        if "metadata" in result:
            metadata = result["metadata"]
            english_metadata = {
                "document_number": metadata.get("so_hieu"),
                "document_type": metadata.get("loai_van_ban"),
                "document_title": metadata.get("ten_van_ban"),
                "issue_date": metadata.get("ngay_ban_hanh_str"),
                "issuing_agency": metadata.get("co_quan_ban_hanh"),
                "effective_date": metadata.get("ngay_hieu_luc_str"),
                "expiry_date": metadata.get("ngay_het_hieu_luc_str"),
                "confidential_level": metadata.get("muc_do_mat", "Công Khai"),
                "source_file": os.path.splitext(filename)[0],
            }

            # Thêm năm ban hành nếu có
            if "nam_ban_hanh" in metadata and metadata["nam_ban_hanh"]:
                try:
                    english_metadata["issue_year"] = int(metadata["nam_ban_hanh"])
                except:
                    english_metadata["issue_year"] = None

            result["metadata"] = english_metadata

        result["success"] = True
        logger.info(
            f"[Comprehensive] Successfully analyzed '{filename}', found {len(result.get('hierarchical_structure', []))} articles")

        return result

    except Exception as e:
        logger.error(f"[Comprehensive] Error analyzing '{filename}': {e}")
        return {
            "metadata": {
                "document_number": None,
                "document_type": None,
                "document_title": None,
                "issue_date": None,
                "issuing_agency": None,
                "effective_date": None,
                "expiry_date": None,
                "confidential_level": "Công Khai",
                "source_file": os.path.splitext(filename)[0],
            },
            "hierarchical_structure": [],
            "success": False,
            "error": str(e)
        }

# Hàm mới thay thế hierarchical_split_law_document() sử dụng phân tích tổng hợp
def process_document_with_comprehensive_analysis(doc_obj: Document) -> List[Document]:
    """
    Hàm mới thay thế hierarchical_split_law_document() sử dụng phân tích tổng hợp
    """
    text = doc_obj.page_content
    source_metadata = doc_obj.metadata.copy()
    filename = source_metadata.get("source_file", source_metadata.get("source", "unknown_file"))

    # Gọi hàm phân tích tổng hợp
    analysis_result = comprehensive_legal_analysis(text, filename)

    if not analysis_result["success"]:
        logger.error(f"Comprehensive analysis failed for {filename}: {analysis_result.get('error')}")
        return []

    # Cập nhật metadata gốc với thông tin đã trích xuất
    updated_metadata = source_metadata.copy()
    updated_metadata.update(analysis_result["metadata"])

    # Suy luận lĩnh vực pháp luật
    updated_metadata["law_field"] = infer_field(text, updated_metadata.get("document_title"))

    final_chunks: List[Document] = []
    hierarchical_structure = analysis_result.get("hierarchical_structure", [])

    for item in hierarchical_structure:
        dieu_code = item.get("dieu_code")
        dieu_title = item.get("dieu_title")
        content = item.get("content")

        if not dieu_code or not content:
            continue

        # Metadata cho chunk này
        block_meta = updated_metadata.copy()
        block_meta["dieu_code"] = dieu_code
        block_meta["dieu_title"] = dieu_title
        block_meta["chunk_title"] = dieu_code + (f" - {dieu_title}" if dieu_title else "")

        # Thông tin đã được LLM trích xuất
        block_meta["penalties"] = item.get("penalties", [])
        block_meta["entity_types"] = item.get("entity_types", [])
        block_meta["cross_references"] = item.get("cross_references", [])

        # Suy luận entity type nếu LLM chưa trích xuất được
        if not block_meta["entity_types"]:
            block_meta["entity_type"] = infer_entity_type(content, updated_metadata.get("law_field"))
        else:
            block_meta["entity_type"] = block_meta["entity_types"]

        # Tạo nội dung chunk
        context_header = f"Excerpt from: {block_meta['chunk_title']}\nIn document: {updated_metadata.get('document_title', filename)}"
        final_page_content = f"Full text: {block_meta['chunk_title']}\n\nContent:\n{content}"

        # Tạo ID cho chunk
        chunk_id = generate_structured_id(
            updated_metadata.get("document_number"),
            [dieu_code],
            filename
        )

        final_chunks.append(Document(
            page_content=final_page_content,
            metadata=block_meta,
            id=chunk_id
        ))

    logger.info(f"[Comprehensive] Generated {len(final_chunks)} chunks for {filename}")
    return final_chunks

#xem lại hàm này để cải tiến sau
def infer_field(text_content: str, doc_title: Optional[str]) -> str:
    """
    CẢI TIẾN V2: Bổ sung từ khóa ngắn gọn, thông tục để xử lý câu hỏi người dùng.
    """
    safe_text_content = str(text_content) if text_content else ""
    safe_doc_title = str(doc_title) if doc_title else ""

    if not safe_doc_title and not safe_text_content:
        return "khac"

    # Chỉ cần một đoạn ngắn để tìm kiếm, giúp tăng hiệu suất
    search_text = (safe_doc_title.lower() + " " + safe_text_content[:2500].lower()).strip()
    title_lower = safe_doc_title.lower()


    field_keywords = {
        # 1. Giao thông
        "giao_thong": [
            ("trật tự, an toàn giao thông đường bộ", 12),
            ("xử phạt vi phạm hành chính trong lĩnh vực giao thông", 11),
            ("giao thông đường bộ", 10),
            ("giao thông đường sắt", 10),
            ("giấy phép lái xe", 9), ("gplx", 9),
            ("đèn tín hiệu giao thông", 8),
            ("vượt đèn đỏ", 9),
            ("nồng độ cồn", 9),
            ("quá tốc độ", 8),
            ("đăng kiểm", 7),
            ("phạt nguội", 7),
            ("xe ô tô", 5), ("ô-tô", 5),
            ("xe máy", 5), ("xe mô tô", 5),
            ("lái xe", 4),
            ("biển báo", 4),
        ],

        # 2. Hình sự
        "hinh_su": [
            ("bộ luật hình sự", 12),
            ("truy cứu trách nhiệm hình sự", 10),
            ("tội phạm", 9),
            ("khởi tố", 8), ("tố tụng hình sự", 8),
            ("điều tra hình sự", 8),
            ("tòa án nhân dân", 7),
            ("viện kiểm sát", 7),
            ("giết người", 9),
            ("cướp giật tài sản", 9),
            ("lừa đảo chiếm đoạt tài sản", 9),
            ("ma túy", 8),
            ("cố ý gây thương tích", 8),
            ("tử hình", 6), ("tù chung thân", 6),
        ],

        # 3. Dân sự
        "dan_su": [
            ("bộ luật dân sự", 12),
            ("bồi thường thiệt hại ngoài hợp đồng", 10),
            ("giao dịch dân sự", 9),
            ("quyền sở hữu", 8),
            ("thừa kế", 9),
            ("di chúc", 9),
            ("hợp đồng dân sự", 7),
            ("tranh chấp dân sự", 6),
            ("nghĩa vụ dân sự", 6),
            ("đại diện, ủy quyền", 5),
        ],

        # 4. Hôn nhân và Gia đình
        "hon_nhan_gia_dinh": [
            ("luật hôn nhân và gia đình", 12),
            ("kết hôn", 9),
            ("ly hôn", 10),
            ("quan hệ giữa vợ và chồng", 8),
            ("tài sản chung của vợ chồng", 8), ("tài sản riêng", 8),
            ("quyền nuôi con", 9),
            ("cấp dưỡng", 8),
            ("mang thai hộ", 7),
            ("giám hộ", 6),
        ],

        # 5. Lao động
        "lao_dong": [
            ("bộ luật lao động", 12),
            ("hợp đồng lao động", 10),
            ("người lao động", 9), ("nlđ", 9),
            ("người sử dụng lao động", 9), ("nsdlđ", 9),
            ("bảo hiểm xã hội", 8), ("bhxh", 8),
            ("tiền lương", 7), ("lương tối thiểu vùng", 7),
            ("thời giờ làm việc", 6), ("thời giờ nghỉ ngơi", 6),
            ("kỷ luật lao động", 6),
            ("sa thải", 7), ("chấm dứt hợp đồng lao động", 7),
            ("an toàn, vệ sinh lao động", 6),
        ],

        # 6. Đất đai
        "dat_dai": [
            ("luật đất đai", 12),
            ("quyền sử dụng đất", 10),
            ("giấy chứng nhận quyền sử dụng đất", 9),
            ("thu hồi đất", 8),
            ("bồi thường, hỗ trợ, tái định cư", 8),
            ("quy hoạch, kế hoạch sử dụng đất", 7),
            ("sổ đỏ", 7), ("sổ hồng", 7),
            ("tranh chấp đất đai", 7),
            ("giá đất", 6),
        ],

        # 7. Doanh nghiệp & Đầu tư
        "doanh_nghiep": [
            ("luật doanh nghiệp", 12),
            ("luật đầu tư", 12),
            ("thành lập doanh nghiệp", 9),
            ("giấy chứng nhận đăng ký doanh nghiệp", 9),
            ("công ty cổ phần", 8),
            ("công ty trách nhiệm hữu hạn", 8),
            ("doanh nghiệp tư nhân", 8),
            ("vốn điều lệ", 7),
            ("cổ đông", 6), ("thành viên góp vốn", 6),
            ("phá sản", 7),
        ],

        # 8. Xây dựng & Nhà ở
        "xay_dung": [
            ("luật xây dựng", 12),
            ("luật nhà ở", 12),
            ("giấy phép xây dựng", 9),
            ("quy hoạch xây dựng", 8),
            ("chủ đầu tư", 7),
            ("dự án đầu tư xây dựng", 7),
            ("hợp đồng xây dựng", 6),
            ("chung cư", 6),
        ],

        # 9. Hành chính
        "hanh_chinh": [
            ("luật xử lý vi phạm hành chính", 10),
            ("xử phạt vi phạm hành chính", 9),
            ("khiếu nại", 8),
            ("tố cáo", 8),
            ("thủ tục hành chính", 7),
            ("công chức", 7), ("viên chức", 7),
            ("cán bộ", 6),
        ],

        # 10. Thuế & Tài chính & Ngân hàng
        "tai_chinh_thue": [
            ("luật quản lý thuế", 12),
            ("luật các tổ chức tín dụng", 12),
            ("thuế giá trị gia tăng", 9), ("thuế gtgt", 9),
            ("thuế thu nhập doanh nghiệp", 9), ("thuế tndn", 9),
            ("thuế thu nhập cá nhân", 9), ("thuế tncn", 9),
            ("ngân sách nhà nước", 8),
            ("hóa đơn điện tử", 7),
            ("kế toán, kiểm toán", 7),
            ("ngân hàng", 6),
            ("trái phiếu", 6),
        ],

        # 11. Môi trường
        "moi_truong": [
            ("luật bảo vệ môi trường", 12),
            ("đánh giá tác động môi trường", 9), ("đtm", 9),
            ("ô nhiễm môi trường", 8),
            ("chất thải", 7), ("chất thải nguy hại", 7),
            ("tài nguyên nước", 6),
            ("khí thải", 6),
        ],

        # 12. Sở hữu trí tuệ
        "so_huu_tri_tue": [
            ("luật sở hữu trí tuệ", 12),
            ("quyền tác giả", 9),
            ("bản quyền", 8),
            ("quyền liên quan", 8),
            ("sáng chế", 8),
            ("nhãn hiệu", 8),
            ("chỉ dẫn địa lý", 7),
        ],

        # 13. Giáo dục
        "giao_duc": [
            ("luật giáo dục", 12),
            ("học sinh, sinh viên", 7),
            ("cơ sở giáo dục", 7),
            ("học phí", 7),
            ("tuyển sinh", 6),
            ("giáo viên", 6),
            ("bằng cấp, chứng chỉ", 5),
        ],

        # 14. Y tế
        "y_te": [
            ("luật khám bệnh, chữa bệnh", 12),
            ("bảo hiểm y tế", 10), ("bhyt", 10),
            ("dược", 8), ("thuốc", 7),
            ("trang thiết bị y tế", 7),
            ("bệnh viện", 6),
            ("an toàn thực phẩm", 6),
        ],
    }

    # ... (phần logic tính điểm và trả về giữ nguyên) ...
    field_scores = {field: 0 for field in field_keywords.keys()}
    for field, weighted_keywords in field_keywords.items():
        score = 0
        for keyword, weight in weighted_keywords:
            if doc_title and keyword in title_lower:
                score += weight * 3
            occurrences_in_text = search_text.count(keyword)
            if occurrences_in_text > 0:
                score += weight * occurrences_in_text
        field_scores[field] = score

    positive_scores = {f: s for f, s in field_scores.items() if s > 0}
    if not positive_scores:
        return "khac"

    # In ra điểm số để debug
    logger.debug(f"Field scores for query '{text_content[:50]}...': {positive_scores}")

    best_field = max(positive_scores, key=positive_scores.get)
    return best_field

#xem lại hàm này để cải tiến sau
def infer_entity_type(query_or_text: str, field: Optional[str]) -> Optional[List[str]]:
    """
    CẢI TIẾN V2: Mở rộng từ khóa, xử lý khi không có field và luôn trả về list.
    """
    text_lower = query_or_text.lower()
    entity_definitions = {
        "giao_thong": {
            "xe_oto": {"keywords": ["ô tô", "xe hơi", "xe con", "xe ô-tô"], "priority": 10},
            "xe_may": {"keywords": ["xe máy", "mô tô", "xe gắn máy", "xe 2 bánh"], "priority": 10},
            # Thêm các từ khóa ngắn gọn hơn
            "nguoi_dieu_khien": {"keywords": ["người điều khiển", "lái xe", "tài xế"], "priority": 9},
            "phuong_tien": {"keywords": ["phương tiện", "xe cộ", "xe"], "priority": 5}, # Thêm "xe cộ"
        },
        "hinh_su": {
            "nguoi_pham_toi": {"keywords": ["tội phạm", "bị can", "bị cáo", "người phạm tội", "kẻ gian"], "priority": 10},
            "nan_nhan": {"keywords": ["nạn nhân", "người bị hại"], "priority": 9},
        },
        "lao_dong": {
            "nguoi_lao_dong": {"keywords": ["người lao động", "nhân viên", "công nhân", "nlđ"], "priority": 10},
            "nguoi_su_dung_lao_dong": {"keywords": ["người sử dụng lao động", "công ty", "doanh nghiệp", "nsdlđ"], "priority": 10},
            "hop_dong_lao_dong": {"keywords": ["hợp đồng lao động", "hđlđ"], "priority": 8},
        },
        "khac": {
            "ca_nhan": {"keywords": ["cá nhân", "người", "công dân", "một người"], "priority": 7},
            "to_chuc": {"keywords": ["tổ chức", "cơ quan", "đơn vị", "công ty"], "priority": 7},
        }
    }

    found_entities = []

    # CẢI TIẾN: Nếu có field, chỉ tìm trong field đó. Nếu không, tìm trong tất cả.
    fields_to_check = [field] if field and field in entity_definitions else list(entity_definitions.keys())

    for f in fields_to_check:
        current_field_entities = entity_definitions.get(f, {})
        sorted_entities = sorted(current_field_entities.items(), key=lambda item: item[1]["priority"], reverse=True)
        for entity_type, definition in sorted_entities:
            sorted_keywords = sorted(definition["keywords"], key=len, reverse=True)
            if any(re.search(r"\b" + re.escape(keyword) + r"\b", text_lower) for keyword in sorted_keywords):
                if entity_type not in found_entities:
                    found_entities.append(entity_type)

    if not found_entities:
        return None

    # Luôn trả về một danh sách các entity tìm được
    return found_entities

# xem lại hàm này để chunking và trích xuất bằng llm sau, hàm def parse_law_item_line cần xem xét lại
def parse_law_item_line(line: str) -> Tuple[Optional[str], str, str]:
    """Phân tích cấu trúc dòng một cách mạnh mẽ và có thứ tự."""
    stripped_line = line.strip()
    patterns = [
        ("phan", r"^\s*(PHẦN\s+(?:THỨ\s+[\w\sÀ-Ỹà-ỹ]+|[IVXLCDM]+|CHUNG))\s*?$"),
        ("chuong", r"^\s*(Chương\s+[IVXLCDM\d]+)\s*?$"),
        ("dieu", r"^\s*(Điều\s+\d+[a-z]?)\.?\s*(.*)"),
    ]
    for item_type, pattern_str in patterns:
        match = re.match(pattern_str, stripped_line, re.IGNORECASE)
        if match:
            if item_type in ["phan", "chuong"]:
                return item_type, match.group(1).strip(), ""
            elif item_type == "dieu":
                return item_type, match.group(1).strip(), match.group(2).strip()

    if stripped_line.isupper() and len(stripped_line.split()) > 1 and len(stripped_line) < 150:
        return "title", "", stripped_line
    if m := re.match(r"^\s*(\d+)\.\s+(.*)", stripped_line):
        return "khoan", m.group(1), m.group(2).strip()
    if m := re.match(r"^\s*([a-zđ])\)\s+(.*)", stripped_line):
        return "diem", m.group(1), m.group(2).strip()
    return None, "", stripped_line

# các hàm đằng sau này đều cần xem lại
def _normalize_money(value_str: str) -> Optional[float]:
    if not value_str: return None
    try:
        return float(value_str.replace(".", "").replace(",", "."))
    except ValueError: return None

def _normalize_duration(value_str: str, unit_str: str) -> Optional[Dict[str, Union[int, str]]]:
    if not value_str or not unit_str: return None
    try:
        value = int(value_str)
        unit = unit_str.lower()
        if unit == "tháng": return {"value": value, "unit": "months"}
        if unit == "năm": return {"value": value, "unit": "years"}
        if unit == "ngày": return {"value": value, "unit": "days"}
        return {"value": value, "unit": unit_str}
    except ValueError: return None

def extract_penalties_from_text(text_content: str) -> List[Dict]:
    """
    Trích xuất các loại hình phạt khác nhau từ một đoạn văn bản.
    Cải tiến: Sử dụng set để tránh thêm các hình phạt trùng lặp.
    """
    if not text_content:
        return []

    penalties = []
    found_original_texts = set() # Set để theo dõi các chuỗi đã tìm thấy

    # --- 1. PHẠT TIỀN ---
    # Ưu tiên bắt khoảng (từ... đến...) trước
    fine_range_pattern = r"phạt tiền từ\s*([\d\.,]+)\s*đồng\s*đến\s*([\d\.,]+)\s*đồng"
    for m in re.finditer(fine_range_pattern, text_content, re.IGNORECASE):
        original_text = m.group(0)
        if original_text not in found_original_texts:
            penalties.append({
                "type": "fine",
                "min_amount": _normalize_money(m.group(1)),
                "max_amount": _normalize_money(m.group(2)),
                "currency": "đồng",
                "original_text": original_text
            })
            found_original_texts.add(original_text)

    # Bắt các mức phạt cố định sau
    fine_fixed_pattern = r"\bphạt tiền\s*([\d\.,]+)\s*đồng\b"
    for m in re.finditer(fine_fixed_pattern, text_content, re.IGNORECASE):
        original_text = m.group(0)
        # Kiểm tra xem nó có phải là một phần của một khoảng đã tìm thấy không
        is_part_of_range = any(original_text in found_range for found_range in found_original_texts)
        if not is_part_of_range and original_text not in found_original_texts:
            penalties.append({
                "type": "fine",
                "amount": _normalize_money(m.group(1)),
                "currency": "đồng",
                "original_text": original_text
            })
            found_original_texts.add(original_text)

    # --- 2. HÌNH PHẠT TÙ ---
    prison_range_pattern = r"phạt tù từ\s*(\d+)\s*(tháng|năm|ngày)\s*đến\s*(\d+)\s*(tháng|năm|ngày)"
    for m in re.finditer(prison_range_pattern, text_content, re.IGNORECASE):
        original_text = m.group(0)
        if original_text not in found_original_texts:
            penalties.append({
                "type": "prison",
                "min_duration": _normalize_duration(m.group(1), m.group(2)),
                "max_duration": _normalize_duration(m.group(3), m.group(4)),
                "original_text": original_text
            })
            found_original_texts.add(original_text)

    prison_fixed_pattern = r"\bphạt tù\s*(\d+)\s*(tháng|năm|ngày)\b"
    for m in re.finditer(prison_fixed_pattern, text_content, re.IGNORECASE):
        original_text = m.group(0)
        is_part_of_range = any(original_text in found_range for found_range in found_original_texts)
        if not is_part_of_range and original_text not in found_original_texts:
            penalties.append({
                "type": "prison",
                "duration": _normalize_duration(m.group(1), m.group(2)),
                "original_text": original_text
            })
            found_original_texts.add(original_text)

    # Tù chung thân và tử hình
    special_prison_patterns = {
        "life_imprisonment": r"phạt tù chung thân",
        "death_penalty": r"phạt tử hình"
    }
    for p_type, pattern in special_prison_patterns.items():
        if match := re.search(pattern, text_content, re.IGNORECASE):
            original_text = match.group(0)
            if original_text not in found_original_texts:
                penalties.append({"type": "prison", "duration_type": p_type, "original_text": original_text})
                found_original_texts.add(original_text)

    # --- 3. TƯỚC QUYỀN SỬ DỤNG GIẤY PHÉP ---
    license_revocation_pattern = r"tước quyền sử dụng giấy phép lái xe\s*(?:từ|từ\s*thời\s*hạn)?\s*(\d+)\s*(tháng)\s*đến\s*(\d+)\s*(tháng)"
    for m in re.finditer(license_revocation_pattern, text_content, re.IGNORECASE):
        original_text = m.group(0)
        if original_text not in found_original_texts:
            penalties.append({
                "type": "license_revocation",
                "min_duration": _normalize_duration(m.group(1), m.group(2)),
                "max_duration": _normalize_duration(m.group(3), m.group(4)),
                "original_text": original_text
            })
            found_original_texts.add(original_text)

    # --- 4. CÁC HÌNH PHẠT KHÁC (Cảnh cáo, tịch thu, v.v.) ---
    other_patterns = {
        "warning": r"\bphạt cảnh cáo\b",
        "confiscation_object_vehicle": r"tịch thu tang vật,? phương tiện vi phạm hành chính",
        "deportation": r"\btrục xuất\b"
    }
    for p_type, pattern in other_patterns.items():
        if match := re.search(pattern, text_content, re.IGNORECASE):
            original_text = match.group(0)
            if original_text not in found_original_texts:
                penalties.append({"type": p_type, "original_text": original_text})
                found_original_texts.add(original_text)

    return penalties


# Cập nhật hàm process_single_file() để sử dụng phân tích tổng hợp
def process_single_file_comprehensive(file_path: str) -> List[Document]:
    """
    Phiên bản cập nhật của process_single_file() sử dụng phân tích tổng hợp
    """
    filename = os.path.basename(file_path)
    logger.info(f"--- Starting Comprehensive Processing Pipeline for: {filename} ---")

    try:
        # --- BƯỚC 1: ĐỌC FILE ---
        raw_content = read_file_content(file_path)
        if not raw_content or not raw_content.strip():
            logger.warning(f"File '{filename}' is empty or unreadable.")
            return []

        # --- BƯỚC 2: LÀM SẠCH VĂN BẢN ---
        # Việc làm sạch trước giúp các bước sau hoạt động chính xác hơn
        if not raw_content.strip():
            logger.warning(f"File '{filename}' is empty. Skipping.")
            return []

        cleaned_content = clean_document_text(raw_content)
        if not cleaned_content.strip():
            logger.warning(f"File '{filename}' is empty after cleaning. Skipping.")
            return []
        logger.debug(f"File '{filename}' cleaned successfully.")

        # Tạo document object với metadata cơ bản
        initial_metadata = {"source": filename, "source_file": os.path.splitext(filename)[0]}
        doc_to_process = Document(page_content=cleaned_content, metadata=initial_metadata)

        # Xử lý với phân tích tổng hợp
        chunks_from_file = process_document_with_comprehensive_analysis(doc_to_process)

        if not chunks_from_file:
            logger.warning(f"File '{filename}' did not yield any chunks after comprehensive processing.")
        else:
            logger.info(
                f"✅ Successfully processed '{filename}' with comprehensive analysis, generated {len(chunks_from_file)} chunks.")

        return chunks_from_file

    except Exception as e:
        logger.error(f"❌ Critical error in comprehensive processing of '{filename}': {e}", exc_info=True)
        return []


# Cập nhật hàm load_process_and_split_documents để sử dụng phân tích tổng hợp
def load_process_and_split_documents_comprehensive(folder_path: str) -> List[Document]:
    """
    Phiên bản cập nhật sử dụng comprehensive analysis
    """
    all_final_chunks = []
    if not os.path.isdir(folder_path):
        logger.error(f"Folder '{folder_path}' does not exist.")
        return all_final_chunks

    # List all supported files
    supported_files = [f for f in os.listdir(folder_path) if f.lower().endswith(('.txt', '.docx'))]
    if not supported_files:
        logger.warning(f"Không tìm thấy file .txt hoặc .docx trong '{folder_path}'.")
        return all_final_chunks

    logger.info(f"Found {len(supported_files)} .txt files. Starting comprehensive processing...")

    for filename in tqdm(supported_files, desc="Processing files with comprehensive analysis"):
        file_path = os.path.join(folder_path, filename)
        try:
            chunks_from_file = process_single_file_comprehensive(file_path)
            all_final_chunks.extend(chunks_from_file)
        except Exception as e:
            logger.error(f"Error in comprehensive processing of file '{filename}': {e}", exc_info=True)

    logger.info(f"Comprehensive processing completed: {len(supported_files)} files, {len(all_final_chunks)} total chunks.")
    return all_final_chunks

def read_file_content(file_path: str) -> str:
    """
    Đọc nội dung từ file .txt hoặc .docx.
    Trả về chuỗi văn bản hoặc ném lỗi nếu không đọc được.
    """
    try:
        if file_path.lower().endswith('.txt'):
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        elif file_path.lower().endswith('.docx'):
            doc = DocxDocument(file_path)
            full_text = [para.text for para in doc.paragraphs if para.text.strip()]  # Bỏ qua đoạn rỗng
            # Xử lý bảng
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        full_text.append(row_text)
            # Xử lý header và footer
            for section in doc.sections:
                header = section.header
                footer = section.footer
                for para in header.paragraphs:
                    if para.text.strip():
                        full_text.append(para.text)
                for para in footer.paragraphs:
                    if para.text.strip():
                        full_text.append(para.text)
            return '\n'.join(full_text)
        else:
            raise ValueError(f"Định dạng file không hỗ trợ: {file_path}. Chỉ hỗ trợ .txt và .docx.")
    except Exception as e:
        logger.error(f"Lỗi khi đọc file {file_path}: {e}", exc_info=True)
        raise